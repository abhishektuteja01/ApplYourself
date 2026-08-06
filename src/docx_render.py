"""ATS-clean resume docx generation.

Determinism boundary (R7): no LLM calls.

render_resume uses templated paragraph styles only: the template
(profile/resume_template.docx) defines the visual formatting, and this module
creates paragraphs in the right styles from the markdown. The template MUST
define these 5 named paragraph styles:
  - Resume Name        (the name at top)
  - Resume Section     (SUMMARY / WORK EXPERIENCE / EDUCATION / ...)
  - Resume Job Header  (role / degree / project headers; bold via inline runs)
  - Resume Body        (contact line, summary text, skills lines)
  - Resume Bullet      (bulleted items)
Forbidden in the template (rejected loudly): tables (ATS-unfriendly) and inline
shapes (images, icons).

Cover letters are the opposite operation — preserve the user's design and fill
placeholders — and live in src/docx_cover_letter.py.
"""
from __future__ import annotations

import re
from pathlib import Path
from typing import TypedDict

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn


REQUIRED_STYLES = [
    "Resume Name",
    "Resume Section",
    "Resume Job Header",
    "Resume Body",
    "Resume Bullet",
]

# Required character-type style for clickable [text](url) markdown.
REQUIRED_CHARACTER_STYLE = "Hyperlink"

# Cover-letter template contract (placeholder fill-in, not markdown render --
# the user's own .docx design is preserved untouched; only these placeholder
# paragraphs get their text replaced). {{BODY}} is special: its paragraph's
# formatting is cloned once per body paragraph, then the placeholder itself
# is removed.


class TemplateMissingError(FileNotFoundError):
    """Template file does not exist on disk."""


class TemplateError(ValueError):
    """Template exists but violates the docx_render contract."""


class Run(TypedDict, total=False):
    text: str
    bold: bool
    url: str   # optional; when set, emit as a Word hyperlink (clickable,
               # Hyperlink character style applied for blue+underline)


class Block(TypedDict, total=False):
    type: str   # "name" | "section_header" | "job_header" | "body" | "bullet"
    text: str
    runs: list[Run]


# =====================================================================
# Public API
# =====================================================================

def render_resume(md_content: str, template_path: Path, out_path: Path) -> None:
    """Render resume markdown to an ATS-clean docx using the template.
    Fail loud on missing template / missing required styles / forbidden
    template elements (tables, inline shapes)."""
    if not template_path.exists():
        raise TemplateMissingError(
            f"ERROR: {template_path} missing. Author it in Word with the ATS "
            f"constraints (Calibri, single column, no tables, bold section "
            f"headers, no images/icons)."
        )
    doc = Document(str(template_path))
    _validate_template(doc, template_path)
    _clear_body(doc)
    blocks = parse_resume_md(md_content)
    _write_blocks(doc, blocks)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


def parse_resume_md(md: str) -> list[Block]:
    """Parse the small resume.md markdown subset into typed blocks.

    Supported subset:
      **Name** or # **Name**            -> name (first such line only)
      **SECTION NAME**                  -> section_header (all-caps content)
      **Bold Prefix** rest of line      -> job_header
      * bullet  /  - bullet             -> bullet
      anything else (non-blank)         -> body

    Inline **bold** runs are split for body / bullet / job_header blocks."""
    blocks: list[Block] = []
    seen_name = False
    for raw in md.split("\n"):
        line = raw.strip()
        if not line:
            continue
        if not seen_name:
            m = re.fullmatch(r"#?\s*\*\*(.+?)\*\*", line)
            if m:
                blocks.append({"type": "name", "text": m.group(1).strip()})
                seen_name = True
                continue
        # Bullet
        if line.startswith(("* ", "- ")):
            body = line[2:].strip()
            blocks.append({
                "type": "bullet", "text": body,
                "runs": _split_inline_runs(body),
            })
            continue
        # Section header: a line that is entirely **ALL CAPS WORDS**
        m = re.fullmatch(r"\*\*([A-Z][A-Z0-9 &/]*[A-Z])\*\*", line)
        if m:
            blocks.append({"type": "section_header", "text": m.group(1).strip()})
            continue
        # Lines starting with **...** followed by more content:
        # - Skills lines whose bold prefix ends with ":" route to body
        #   (e.g. "**Programming:** Python, ..." -> Resume Body w/ inline bold).
        # - Everything else routes to job_header (role / degree / project headers).
        m = re.match(r"\*\*(.+?)\*\*(.*)", line)
        if m and m.group(2).strip():
            bold_prefix = m.group(1).strip()
            if bold_prefix.endswith(":"):
                blocks.append({
                    "type": "body", "text": line,
                    "runs": _split_inline_runs(line),
                })
                continue
            runs: list[Run] = [
                {"text": bold_prefix, "bold": True},
            ]
            tail = m.group(2)
            if tail:
                runs.extend(_split_inline_runs(tail))
            blocks.append({"type": "job_header", "text": line, "runs": runs})
            continue
        # Plain body (may contain inline bold)
        blocks.append({
            "type": "body", "text": line,
            "runs": _split_inline_runs(line),
        })
    return blocks


# =====================================================================
# Internals
# =====================================================================

def _validate_template(doc, path: Path) -> None:
    style_names = {s.name for s in doc.styles}
    missing = [s for s in REQUIRED_STYLES if s not in style_names]
    if missing:
        raise TemplateError(
            f"Template {path} missing required paragraph styles: {missing}. "
            f"Each maps to a markdown block type — define them in Word's Styles "
            f"panel. Required set: {REQUIRED_STYLES}"
        )
    if REQUIRED_CHARACTER_STYLE not in style_names:
        raise TemplateError(
            f"Template {path} missing required character style "
            f"{REQUIRED_CHARACTER_STYLE!r}. Word ships this style by default "
            f"-- if it was removed, apply Word's built-in Hyperlink style to "
            f"any character once to recreate it, then re-save the template."
        )
    if len(doc.tables) > 0:
        raise TemplateError(
            f"Template {path} contains {len(doc.tables)} table(s), which are "
            f"ATS-unfriendly. Remove all tables from the template."
        )
    if len(doc.inline_shapes) > 0:
        raise TemplateError(
            f"Template {path} contains {len(doc.inline_shapes)} inline shape(s) "
            f"(images/icons), which are ATS-unfriendly. Remove from the template."
        )


def _clear_body(doc) -> None:
    """Remove all paragraphs from the document body. Styles are preserved."""
    for p in list(doc.paragraphs):
        elem = p._element
        elem.getparent().remove(elem)


def _write_blocks(doc, blocks: list[Block]) -> None:
    """Write blocks to the doc body.

    Two layout rules applied here (kept here vs in parse_resume_md so the
    markdown shape stays inspection-friendly):

    1. The first body block immediately after the name block is the contact
       line. Force-center it via WD_ALIGN_PARAGRAPH.CENTER on just that
       paragraph; the style stays Resume Body so the font / size inherit.
       Per the template contract (5 styles only), we do not add a sixth
       Resume Contact style.
    2. Any run text containing a tab character is split at the tab;
       add_run("\\t") is emitted between the pieces so Word's layout engine
       advances to the paragraph style's next tab stop. Resume Job Header
       carries a right tab stop at 19.05 cm -- this is what right-aligns
       the date on "**Title** ... <TAB> Date" job-header lines.
    """
    name_just_emitted = False
    for block in blocks:
        btype = block["type"]
        if btype == "name":
            p = doc.add_paragraph(style="Resume Name")
            p.add_run(block["text"])
            name_just_emitted = True
            continue
        if btype == "section_header":
            p = doc.add_paragraph(style="Resume Section")
            p.add_run(block["text"])
            name_just_emitted = False
            continue
        style_map = {
            "job_header": "Resume Job Header",
            "body":       "Resume Body",
            "bullet":     "Resume Bullet",
        }
        if btype not in style_map:
            raise ValueError(f"Unknown block type: {btype}")
        p = doc.add_paragraph(style=style_map[btype])
        is_contact = btype == "body" and name_just_emitted
        if is_contact:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        runs = block.get("runs") or [{"text": block["text"], "bold": False}]
        for r in runs:
            _emit_run_splitting_tabs(
                p, r["text"], bold=bool(r.get("bold")), url=r.get("url"),
            )
        name_just_emitted = False


def _emit_run_splitting_tabs(paragraph, text: str, *, bold: bool,
                              url: str | None = None) -> None:
    """Add runs to paragraph, splitting on tab characters and routing
    URL-bearing runs through the Word hyperlink path. Each "\\t" becomes
    its own non-bold non-link run so Word's layout engine advances to the
    next tab stop. Tabs inside a URL-bearing run are silently dropped --
    URLs don't sensibly contain tab stops."""
    if not text:
        return
    if url:
        # Hyperlinks are atomic in Word: a single <w:hyperlink> element with
        # a single visible run. Tabs inside link text don't make sense in
        # any resume / outreach context we care about.
        _add_hyperlink(paragraph, text, url)
        return
    if "\t" not in text:
        run = paragraph.add_run(text)
        if bold:
            run.bold = True
        return
    parts = text.split("\t")
    for i, part in enumerate(parts):
        if i > 0:
            paragraph.add_run("\t")
        if part:
            run = paragraph.add_run(part)
            if bold:
                run.bold = True


def _add_hyperlink(paragraph, text: str, url: str) -> None:
    """Insert a real Word hyperlink into the paragraph. Uses the template's
    built-in 'Hyperlink' character style for blue+underline rendering."""
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), REQUIRED_CHARACTER_STYLE)
    rPr.append(rStyle)
    new_run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    t.set(qn("xml:space"), "preserve")
    new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


_BOLD_RE = re.compile(r"\*\*(.+?)\*\*")
_LINK_RE = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")


def _split_inline_runs(text: str) -> list[Run]:
    """Split text into runs at **bold** AND [text](url) boundaries.

    Name is historical -- handles bold and links now. Plain text becomes
    plain runs; **x** becomes a bold run; [t](u) becomes a link run (bold
    inside a link or a link inside bold are NOT supported in v1 -- if a
    bold span and a link span overlap, whichever started earlier wins and
    the later overlapping match is dropped)."""
    matches: list[tuple[int, int, Run]] = []
    for m in _BOLD_RE.finditer(text):
        matches.append((m.start(), m.end(), {"text": m.group(1), "bold": True}))
    for m in _LINK_RE.finditer(text):
        matches.append((m.start(), m.end(),
                        {"text": m.group(1), "bold": False, "url": m.group(2)}))
    matches.sort(key=lambda x: x[0])

    runs: list[Run] = []
    pos = 0
    for start, end, run in matches:
        if start < pos:
            continue   # overlap with a prior match; drop
        if start > pos:
            runs.append({"text": text[pos:start], "bold": False})
        runs.append(run)
        pos = end
    if pos < len(text):
        runs.append({"text": text[pos:], "bold": False})
    if not runs:
        runs.append({"text": text, "bold": False})
    return runs
