"""Dump a resume's text so /onboarding can read it.

    uv run profile-extract <file>

Deterministic (R7): plain text extraction, no judgment, no LLM. The onboarding
session reads the output and interviews the user about it; nothing here decides
what is a bullet.

.docx is the case that needs code -- Claude cannot read a Word file directly.
.md and .txt pass through so one command works for every input. .pdf is handled
natively by the Read tool, so this refuses it with a pointer rather than pulling
in a PDF dependency the rest of the pipeline does not need.

Extraction covers three places resumes hide text:
  - body paragraphs, in document order
  - table cells, in document order with the paragraphs (many resume templates
    are a single invisible table, and a paragraph-only walk returns nothing)
  - headers and footers, where contact details often live
"""
from __future__ import annotations

import argparse
import sys
from pathlib import Path

from docx import Document
from docx.table import Table
from docx.text.paragraph import Paragraph

PASSTHROUGH_SUFFIXES = {".md", ".txt", ".markdown"}
DOCX_SUFFIXES = {".docx"}
# .doc is the pre-2007 binary format; python-docx cannot read it at all.
UNSUPPORTED = {
    ".pdf": (
        "PDF is read natively by Claude's Read tool — point it at the file "
        "directly instead of using this command."
    ),
    ".doc": (
        "legacy .doc is not readable by python-docx. Open it in Word and save "
        "as .docx, or export to PDF."
    ),
    ".pages": (
        "Pages documents are not readable here. Export to .docx or PDF first."
    ),
}


def _iter_block_items(parent):
    """Yield Paragraph and Table objects in document order.

    python-docx exposes .paragraphs and .tables as separate flat lists, which
    loses the interleaving and silently drops table content from a linear read.
    """
    body = parent.element.body
    for child in body.iterchildren():
        if child.tag.endswith("}p"):
            yield Paragraph(child, parent)
        elif child.tag.endswith("}tbl"):
            yield Table(child, parent)


def _table_lines(table: Table) -> list[str]:
    lines = []
    for row in table.rows:
        cells = [c.text.strip().replace("\n", " ") for c in row.cells]
        # Collapse the duplicates a merged cell produces across a row.
        deduped = [c for i, c in enumerate(cells) if c and (i == 0 or c != cells[i - 1])]
        if deduped:
            lines.append(" | ".join(deduped))
    return lines


def _section_lines(doc) -> list[str]:
    """Header and footer text, which render_resume preserves and readers miss."""
    lines: list[str] = []
    for i, section in enumerate(doc.sections):
        for label, part in (("header", section.header), ("footer", section.footer)):
            texts = [p.text.strip() for p in part.paragraphs if p.text.strip()]
            for table in part.tables:
                texts.extend(_table_lines(table))
            if texts:
                lines.append(f"[section {i + 1} {label}]")
                lines.extend(texts)
    return lines


def extract_docx(path: Path) -> str:
    doc = Document(str(path))
    lines: list[str] = []
    for block in _iter_block_items(doc):
        if isinstance(block, Paragraph):
            text = block.text.strip()
            if text:
                lines.append(text)
        else:
            lines.extend(_table_lines(block))
    section_lines = _section_lines(doc)
    if section_lines:
        lines.append("")
        lines.extend(section_lines)
    return "\n".join(lines)


def extract(path: Path) -> str:
    """Plain text for a resume file. Raises ValueError with a usable message."""
    if not path.exists():
        raise ValueError(f"{path} does not exist")
    suffix = path.suffix.lower()
    if suffix in UNSUPPORTED:
        raise ValueError(f"{path.name}: {UNSUPPORTED[suffix]}")
    if suffix in PASSTHROUGH_SUFFIXES:
        return path.read_text(encoding="utf-8")
    if suffix in DOCX_SUFFIXES:
        return extract_docx(path)
    raise ValueError(
        f"{path.name}: unsupported extension {suffix!r}. Supported: "
        f"{', '.join(sorted(DOCX_SUFFIXES | PASSTHROUGH_SUFFIXES))}."
    )


def main() -> int:
    parser = argparse.ArgumentParser(
        prog="profile-extract",
        description="Dump a resume's text (.docx, .md, .txt) to stdout.",
    )
    parser.add_argument("file", type=Path, help="resume file to extract")
    args = parser.parse_args()
    try:
        text = extract(args.file)
    except ValueError as exc:
        print(f"ERROR: {exc}", file=sys.stderr)
        return 1
    if not text.strip():
        print(
            f"ERROR: {args.file} produced no text. If it is a scanned image, "
            f"there is nothing to extract.",
            file=sys.stderr,
        )
        return 1
    print(text)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
