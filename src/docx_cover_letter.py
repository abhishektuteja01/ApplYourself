"""Cover-letter rendering: fill placeholders in the user's own .docx design.

This renderer mutates an existing template in place and must leave every
paragraph, style, header, footer and image it does not own byte-for-byte
untouched. The resume renderer next door does the opposite — it clears the body
and writes fresh blocks — so the two share nothing but the template contract.
"""
from __future__ import annotations

import copy
from pathlib import Path

from docx import Document
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

from src.docx_render import TemplateError, TemplateMissingError

# The placeholder contract this renderer fills.
COVER_LETTER_REQUIRED_PLACEHOLDERS = ["{{SALUTATION}}", "{{BODY}}"]
COVER_LETTER_OPTIONAL_PLACEHOLDERS = ["{{DATE}}", "{{CLOSING}}", "{{SIGNOFF_NAME}}"]

def render_cover_letter(content: dict, template_path: Path, out_path: Path) -> None:
    """Fill placeholder paragraphs in the user's own cover-letter .docx
    design, leaving every other paragraph/style/header/footer/image in
    that template byte-for-byte untouched.

    content keys:
      salutation   (str, required)  -- e.g. "Dear Hiring Manager,"
      body         (list[str], required) -- one entry per body paragraph
      date         (str, optional)
      closing      (str, optional)  -- e.g. "Sincerely,"
      signoff_name (str, optional)  -- e.g. "Jane Doe"

    Template contract: somewhere in the template body, a paragraph whose
    stripped text is EXACTLY one of COVER_LETTER_REQUIRED_PLACEHOLDERS /
    COVER_LETTER_OPTIONAL_PLACEHOLDERS marks where that content goes.
    {{SALUTATION}}/{{DATE}}/{{CLOSING}}/{{SIGNOFF_NAME}} are replaced
    in place (the paragraph's first run keeps its formatting). {{BODY}}'s
    paragraph is cloned once per entry in content['body'] (preserving its
    style/run formatting), then the placeholder paragraph is removed.
    Optional placeholders absent from the template are simply skipped --
    the template already has static text there."""
    if not template_path.exists():
        raise TemplateMissingError(
            f"ERROR: {template_path} missing. Save your cover letter design "
            f"there with placeholder paragraphs: {COVER_LETTER_REQUIRED_PLACEHOLDERS} "
            f"required, {COVER_LETTER_OPTIONAL_PLACEHOLDERS} optional."
        )
    doc = Document(str(template_path))
    _fill_cover_letter_placeholders(doc, content, template_path)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


def _paragraph_full_text(paragraph) -> str:
    """All visible text in a paragraph, including runs nested inside
    w:hyperlink and w:sdt (Word 'content control' -- the wrapper Word's
    built-in templates use for click-to-type placeholders). Plain
    paragraph.text only walks direct w:r children and misses both, which
    silently breaks placeholder detection against real Word templates."""
    return "".join(t.text or "" for t in paragraph._p.iter(qn("w:t")))


def _iter_paragraphs(doc):
    """Every body paragraph, including those inside table cells (nested tables
    included). doc.paragraphs walks only direct body children, and many Word
    letter designs lay the whole page out in an invisible table -- placeholders
    in those cells would be invisible to detection and to filling."""
    yield from doc.paragraphs
    tables = list(doc.tables)
    while tables:
        table = tables.pop()
        for row in table.rows:
            for cell in row.cells:
                yield from cell.paragraphs
                tables.extend(cell.tables)


def list_cover_letter_placeholders(template_path: Path) -> set[str]:
    """Return which COVER_LETTER_REQUIRED_PLACEHOLDERS /
    COVER_LETTER_OPTIONAL_PLACEHOLDERS tokens are present in
    template_path (using _paragraph_full_text, so w:hyperlink/w:sdt
    nesting is accounted for). /cover-letter uses this to decide which
    optional fields (date/closing/signoff_name) are worth generating --
    callers must NOT re-implement detection with plain paragraph.text,
    which silently misses tokens nested in Word content controls."""
    doc = Document(str(template_path))
    all_tokens = set(COVER_LETTER_REQUIRED_PLACEHOLDERS) | set(COVER_LETTER_OPTIONAL_PLACEHOLDERS)
    return {
        text for p in _iter_paragraphs(doc)
        if (text := _paragraph_full_text(p).strip()) in all_tokens
    }


def _fill_cover_letter_placeholders(doc, content: dict, path: Path) -> None:
    found: dict[str, object] = {}
    counts: dict[str, int] = {}
    all_tokens = set(COVER_LETTER_REQUIRED_PLACEHOLDERS) | set(COVER_LETTER_OPTIONAL_PLACEHOLDERS)
    for p in _iter_paragraphs(doc):
        text = _paragraph_full_text(p).strip()
        if text in all_tokens:
            counts[text] = counts.get(text, 0) + 1
            found[text] = p

    duplicates = [t for t, n in counts.items() if n > 1]
    if duplicates:
        raise TemplateError(
            f"Template {path} has more than one paragraph containing "
            f"{duplicates}. Each placeholder token must appear exactly once."
        )

    missing = [ph for ph in COVER_LETTER_REQUIRED_PLACEHOLDERS if ph not in found]
    if missing:
        raise TemplateError(
            f"Template {path} missing required placeholder paragraph(s) "
            f"{missing}. Add a paragraph containing EXACTLY that token "
            f"(on its own line, with whatever formatting you want) "
            f"wherever the generated content should go."
        )

    if not content.get("salutation"):
        raise ValueError("content['salutation'] is required and must be non-empty.")
    if not content.get("body"):
        raise ValueError("content['body'] is required and must be a non-empty list of paragraph strings.")

    simple_values = {
        "{{DATE}}": content.get("date"),
        "{{SALUTATION}}": content.get("salutation"),
        "{{CLOSING}}": content.get("closing"),
        "{{SIGNOFF_NAME}}": content.get("signoff_name"),
    }
    for token, value in simple_values.items():
        if token in found:
            # Found-in-template placeholders ALWAYS get filled, even with ""
            # for an unsupplied optional field -- leaving the raw {{TOKEN}}
            # text in a generated letter is a worse failure than a blank line.
            _set_paragraph_text(found[token], value or "")

    _expand_body_placeholder(found["{{BODY}}"], content["body"])


def _set_paragraph_text(paragraph, text: str) -> None:
    """Replace a placeholder paragraph's visible text in place, finding
    the underlying w:t regardless of w:hyperlink/w:sdt nesting (see
    _paragraph_full_text). The first w:t's run keeps its character
    formatting; if that run sits inside a Word content control (w:sdt),
    the control is unwrapped (replaced by the bare run) so no stale
    'click here to type' plumbing survives into the generated letter.
    Any other w:t left in the paragraph (e.g. a split token) is cleared."""
    p_elem = paragraph._p
    t_elems = list(p_elem.iter(qn("w:t")))
    if not t_elems:
        paragraph.add_run(text)
        return
    first_t = t_elems[0]
    run_elem = first_t.getparent()
    _unwrap_run_ancestor(run_elem, p_elem)
    first_t.text = text
    for t in t_elems[1:]:
        t.text = ""


_UNWRAP_TAGS = (qn("w:sdt"), qn("w:hyperlink"))


def _unwrap_run_ancestor(run_elem, paragraph_elem) -> bool:
    """If run_elem sits inside a Word content control (w:sdt) or a hyperlink
    (w:hyperlink) within paragraph_elem, replace that wrapper with run_elem
    itself (preserving run_elem's formatting/position) and return True. No-op
    (returns False) for a direct child. Both wrappers must go: an sdt leaves
    stale 'click here to type' plumbing in the letter, and a hyperlink would
    render the generated text as a live link to the template's URL."""
    node = run_elem
    parent = node.getparent()
    while parent is not None and parent is not paragraph_elem:
        if parent.tag in _UNWRAP_TAGS:
            parent.getparent().replace(parent, run_elem)
            return True
        node = parent
        parent = node.getparent()
    return False


def _expand_body_placeholder(placeholder_p, paragraphs: list[str]) -> None:
    """Clone the {{BODY}} paragraph's XML once per entry in `paragraphs`
    (preserving its style/run formatting), insert each clone immediately
    before the placeholder, set its text, then remove the now-empty
    placeholder paragraph itself."""
    anchor = placeholder_p._p
    for text in paragraphs:
        new_elem = copy.deepcopy(anchor)
        anchor.addprevious(new_elem)
        new_paragraph = Paragraph(new_elem, placeholder_p._parent)
        _set_paragraph_text(new_paragraph, text)
    anchor.getparent().remove(anchor)
