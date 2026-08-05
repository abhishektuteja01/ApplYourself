"""Tests for src/docx_render.py — ATS-clean docx generation.

Uses a programmatic fixture template at test setup so tests don't depend on
the user's hand-authored template (which is part of slice 3 / onboarding)."""
from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document
from docx.enum.style import WD_STYLE_TYPE

from src.docx_render import (
    REQUIRED_STYLES,
    TemplateError,
    TemplateMissingError,
    parse_resume_md,
    render_cover_letter,
    render_resume,
)


# ---------- fixture templates ----------

def _make_template(path: Path, *, styles: list[str] | None = None,
                   add_table: bool = False,
                   include_hyperlink_style: bool = True) -> Path:
    """Build a minimal template with the requested styles. styles=None means
    all 5 required paragraph styles. The Hyperlink character style is added
    explicitly because python-docx's default Document() leaves it latent
    (not realised in styles.xml) until first use."""
    if styles is None:
        styles = list(REQUIRED_STYLES)
    doc = Document()
    existing = {s.name for s in doc.styles}
    for name in styles:
        if name not in existing:
            doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    if include_hyperlink_style and "Hyperlink" not in {s.name for s in doc.styles}:
        doc.styles.add_style("Hyperlink", WD_STYLE_TYPE.CHARACTER)
    if add_table:
        doc.add_table(rows=2, cols=2)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))
    return path


# ---------- fail-loud paths ----------

def test_render_fails_loud_on_missing_template(tmp_path):
    out = tmp_path / "resume.docx"
    nope = tmp_path / "does_not_exist.docx"
    with pytest.raises(TemplateMissingError) as exc:
        render_resume("**Name**", nope, out)
    msg = str(exc.value)
    # Must be actionable: name the file + the ATS constraints to author
    assert "does_not_exist.docx missing" in msg
    assert "ATS" in msg
    assert "Calibri" in msg
    assert not out.exists()


def test_render_fails_loud_on_missing_style(tmp_path):
    # Build a template missing "Resume Bullet"
    template = _make_template(
        tmp_path / "tmpl.docx",
        styles=[s for s in REQUIRED_STYLES if s != "Resume Bullet"],
    )
    out = tmp_path / "resume.docx"
    with pytest.raises(TemplateError) as exc:
        render_resume("**Name**", template, out)
    assert "Resume Bullet" in str(exc.value)
    assert not out.exists()


def test_render_fails_loud_on_tables_in_template(tmp_path):
    template = _make_template(tmp_path / "tmpl.docx", add_table=True)
    out = tmp_path / "resume.docx"
    with pytest.raises(TemplateError) as exc:
        render_resume("**Name**", template, out)
    assert "table" in str(exc.value).lower()
    assert not out.exists()


# ---------- render produces valid docx ----------

def test_render_creates_non_empty_docx(tmp_path):
    template = _make_template(tmp_path / "tmpl.docx")
    out = tmp_path / "resume.docx"
    md = """**Jane Doe**

contact line

**SUMMARY**

A short summary paragraph.

**WORK EXPERIENCE**

**Advisory Analyst** - Acme Corp | May 2022 – Jul 2024

* Owned the daily report.
"""
    render_resume(md, template, out)
    assert out.exists()
    assert out.stat().st_size > 0
    doc = Document(str(out))
    assert len(doc.paragraphs) > 0


def test_render_output_has_no_tables(tmp_path):
    template = _make_template(tmp_path / "tmpl.docx")
    out = tmp_path / "resume.docx"
    render_resume("**Name**\n\n* bullet", template, out)
    doc = Document(str(out))
    assert len(doc.tables) == 0
    assert len(doc.inline_shapes) == 0


def test_render_applies_required_styles_by_name(tmp_path):
    template = _make_template(tmp_path / "tmpl.docx")
    out = tmp_path / "resume.docx"
    md = """**Jane**

**SUMMARY**

A line.

* a bullet
"""
    render_resume(md, template, out)
    doc = Document(str(out))
    style_names = [p.style.name for p in doc.paragraphs]
    assert "Resume Name" in style_names
    assert "Resume Section" in style_names
    assert "Resume Body" in style_names
    assert "Resume Bullet" in style_names


# ---------- parse_resume_md block types ----------

def test_parse_resume_md_name_block_only_first():
    blocks = parse_resume_md("**Jane**\n\n**Other Bold**")
    types = [b["type"] for b in blocks]
    # First **bold** -> name; second is a single bold line treated as job_header? No,
    # a single bold word with no trailing content isn't job_header; falls through to body.
    assert types[0] == "name"
    # The second bold line has no following content, so it is either section_header
    # (if all caps) or body. "Other Bold" is mixed case -> body.
    assert types[1] in {"body", "job_header"}


def test_parse_resume_md_section_header_all_caps():
    blocks = parse_resume_md("**Jane**\n\n**WORK EXPERIENCE**\n")
    section = [b for b in blocks if b["type"] == "section_header"]
    assert len(section) == 1
    assert section[0]["text"] == "WORK EXPERIENCE"


def test_parse_resume_md_job_header_with_bold_prefix():
    blocks = parse_resume_md(
        "**Jane**\n\n**Advisory Analyst** - Acme Corp | May 2022 – Jul 2024"
    )
    job = [b for b in blocks if b["type"] == "job_header"]
    assert len(job) == 1
    runs = job[0]["runs"]
    assert runs[0]["bold"] is True
    assert runs[0]["text"] == "Advisory Analyst"
    # Some plain-text run follows
    assert any(not r["bold"] and "Acme Corp" in r["text"] for r in runs)


def test_parse_resume_md_bullet():
    blocks = parse_resume_md("**N**\n\n* Owned the daily report.")
    bullets = [b for b in blocks if b["type"] == "bullet"]
    assert len(bullets) == 1
    assert bullets[0]["text"] == "Owned the daily report."


def test_parse_resume_md_bullet_with_inline_bold():
    blocks = parse_resume_md("**N**\n\n* **Programming:** Python and SQL")
    bullets = [b for b in blocks if b["type"] == "bullet"]
    runs = bullets[0]["runs"]
    bold_runs = [r for r in runs if r["bold"]]
    assert any(r["text"] == "Programming:" for r in bold_runs)


def test_parse_resume_md_skills_line_colon_prefix_routes_to_body():
    """**Prefix:** ... routes to body with inline bold runs, NOT job_header."""
    blocks = parse_resume_md("**N**\n\n**Programming:** Python, SQL, Java")
    headers = [b for b in blocks if b["type"] == "job_header"]
    bodies = [b for b in blocks if b["type"] == "body"]
    assert len(headers) == 0
    assert len(bodies) == 1
    runs = bodies[0]["runs"]
    bold_prefix = [r for r in runs if r["bold"] and r["text"] == "Programming:"]
    plain_tail = [r for r in runs if not r["bold"] and "Python" in r["text"]]
    assert bold_prefix and plain_tail


def test_parse_resume_md_job_header_without_trailing_colon_stays_job_header():
    """Sanity: prefixes that do NOT end with ':' (job titles, degree titles)
    keep routing to job_header so the right-tab-stop date logic still applies."""
    blocks = parse_resume_md(
        "**N**\n\n**Advisory Analyst** - Acme Corp\tMay 2022 – Jul 2024"
    )
    headers = [b for b in blocks if b["type"] == "job_header"]
    assert len(headers) == 1
    assert headers[0]["runs"][0]["text"] == "Advisory Analyst"


def test_render_emits_tab_in_job_header_paragraph(tmp_path):
    """A literal \\t in a job_header line must reach the docx as a tab
    character so the template's right tab stop catches the date."""
    template = _make_template(tmp_path / "tmpl.docx")
    out = tmp_path / "resume.docx"
    md = "**Jane**\n\n**Advisory Analyst** - Acme Corp\tMay 2022 – Jul 2024"
    render_resume(md, template, out)
    doc = Document(str(out))
    # Find the job_header paragraph
    headers = [p for p in doc.paragraphs if p.style.name == "Resume Job Header"]
    assert len(headers) == 1
    # The full text of the paragraph should contain a tab character
    assert "\t" in headers[0].text
    # And it should split company-on-left from date-on-right
    left, right = headers[0].text.split("\t", 1)
    assert "Acme Corp" in left
    assert "May 2022" in right


def test_render_centers_contact_line_first_body_after_name(tmp_path):
    """The first body block immediately after the name block is the contact
    line; render must force-center just that paragraph (Resume Body style
    still, alignment override)."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    template = _make_template(tmp_path / "tmpl.docx")
    out = tmp_path / "resume.docx"
    md = "**Jane**\n\ncontact line\n\nA later body line."
    render_resume(md, template, out)
    doc = Document(str(out))
    body_paras = [p for p in doc.paragraphs if p.style.name == "Resume Body"]
    assert len(body_paras) == 2
    # First body after name = contact line = centered
    assert body_paras[0].alignment == WD_ALIGN_PARAGRAPH.CENTER
    # Second body = NOT forced-centered (inherits style default = None / left)
    assert body_paras[1].alignment != WD_ALIGN_PARAGRAPH.CENTER


def test_render_does_not_center_when_section_header_intervenes(tmp_path):
    """If something other than body comes right after name, no contact
    centering applies to the next body line."""
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    template = _make_template(tmp_path / "tmpl.docx")
    out = tmp_path / "resume.docx"
    md = "**Jane**\n\n**SUMMARY**\n\nA summary line."
    render_resume(md, template, out)
    doc = Document(str(out))
    body_paras = [p for p in doc.paragraphs if p.style.name == "Resume Body"]
    assert len(body_paras) == 1
    # The summary line is the first body, but it's NOT immediately after name
    # (the SUMMARY section header intervenes), so no centering forced.
    assert body_paras[0].alignment != WD_ALIGN_PARAGRAPH.CENTER


# ---------- inline links ----------

def test_parse_resume_md_inline_link_emits_url_run():
    """[text](url) produces a run with url field set."""
    blocks = parse_resume_md("**N**\n\nVisit [Portfolio](https://example.com) today")
    bodies = [b for b in blocks if b["type"] == "body"]
    assert len(bodies) == 1
    runs = bodies[0]["runs"]
    link_runs = [r for r in runs if r.get("url")]
    assert len(link_runs) == 1
    assert link_runs[0]["text"] == "Portfolio"
    assert link_runs[0]["url"] == "https://example.com"


def test_parse_resume_md_mailto_link():
    """[email](mailto:email) is recognised the same as http links."""
    blocks = parse_resume_md(
        "**N**\n\n[me@example.com](mailto:me@example.com)"
    )
    bodies = [b for b in blocks if b["type"] == "body"]
    runs = bodies[0]["runs"]
    link_runs = [r for r in runs if r.get("url")]
    assert len(link_runs) == 1
    assert link_runs[0]["url"] == "mailto:me@example.com"


def test_parse_resume_md_link_alongside_bold_in_body():
    """A body line containing both **bold** and [link](url) splits into
    plain + bold + plain + link + plain runs.
    (A line that STARTS with **bold** routes to job_header by the existing
    rule -- this test deliberately doesn't start with bold so it stays body.)"""
    blocks = parse_resume_md(
        "**N**\n\nThe **bold** part and a [link](https://x) here"
    )
    bodies = [b for b in blocks if b["type"] == "body"]
    assert len(bodies) == 1
    runs = bodies[0]["runs"]
    assert any(r.get("bold") and r["text"] == "bold" for r in runs)
    assert any(r.get("url") == "https://x" and r["text"] == "link" for r in runs)
    assert any("here" in r["text"] and not r.get("bold") and not r.get("url")
               for r in runs)


def test_render_emits_hyperlink_element_with_correct_target(tmp_path):
    """Rendered docx contains a w:hyperlink element pointing at the URL."""
    template = _make_template(tmp_path / "tmpl.docx")
    out = tmp_path / "resume.docx"
    render_resume(
        "**N**\n\n[Portfolio](https://janedoe.example.com/) and [Email](mailto:x@y.com)",
        template, out,
    )
    doc = Document(str(out))
    W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    hlinks = doc.element.body.findall(f".//{W_NS}hyperlink")
    assert len(hlinks) == 2
    # Verify both target URLs landed in the doc's relationships
    targets = [r.target_ref for r in doc.part.rels.values()
               if "hyperlink" in r.reltype]
    assert any("janedoe.example.com" in t for t in targets)
    assert any("mailto:x@y.com" == t for t in targets)


def test_render_emits_hyperlink_with_hyperlink_char_style(tmp_path):
    """Each hyperlink's inner run carries the 'Hyperlink' character style
    so Word renders it blue+underlined per the template's style def."""
    template = _make_template(tmp_path / "tmpl.docx")
    out = tmp_path / "resume.docx"
    render_resume("**N**\n\n[t](https://x)", template, out)
    doc = Document(str(out))
    W_NS = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    rStyles = doc.element.body.findall(f".//{W_NS}hyperlink//{W_NS}rStyle")
    assert len(rStyles) == 1
    assert rStyles[0].get(f"{W_NS}val") == "Hyperlink"


def test_validate_template_requires_hyperlink_character_style(tmp_path):
    """Template missing the Hyperlink character style fails loud."""
    from src.docx_render import _validate_template
    # Default python-docx Document() has Hyperlink. Delete it then validate.
    template = _make_template(tmp_path / "tmpl.docx")
    doc = Document(str(template))
    for style in list(doc.styles):
        if style.name == "Hyperlink":
            style.delete()
    with pytest.raises(TemplateError, match="Hyperlink"):
        _validate_template(doc, template)


# =====================================================================
# render_cover_letter — placeholder fill-in over the user's own design
# =====================================================================

def _make_cl_template(path: Path, *, tokens=("{{SALUTATION}}", "{{BODY}}"),
                      duplicate: str | None = None,
                      static_text=("Static header line",),
                      body_style: str | None = None,
                      body_bold: bool = False) -> Path:
    """A cover-letter template: static paragraphs plus one paragraph per token.
    body_style/body_bold mark the {{BODY}} paragraph so clone-formatting is
    observable in the output."""
    doc = Document()
    for line in static_text:
        doc.add_paragraph(line)
    for tok in tokens:
        if tok == "{{BODY}}" and (body_style or body_bold):
            p = doc.add_paragraph()
            if body_style:
                if body_style not in {s.name for s in doc.styles}:
                    doc.styles.add_style(body_style, WD_STYLE_TYPE.PARAGRAPH)
                p.style = doc.styles[body_style]
            run = p.add_run(tok)
            run.bold = body_bold
        else:
            doc.add_paragraph(tok)
    if duplicate:
        doc.add_paragraph(duplicate)
    path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(path))
    return path


def _texts(path: Path) -> list[str]:
    return [p.text for p in Document(str(path)).paragraphs]


def _content(**over) -> dict:
    c = {"salutation": "Dear Hiring Manager,", "body": ["First para.", "Second para."]}
    c.update(over)
    return c


# ---------- fail-loud paths ----------

def test_cover_letter_fails_loud_on_missing_template(tmp_path):
    out = tmp_path / "cl.docx"
    with pytest.raises(TemplateMissingError) as exc:
        render_cover_letter(_content(), tmp_path / "nope.docx", out)
    msg = str(exc.value)
    assert "nope.docx missing" in msg
    # actionable: names both placeholder sets
    assert "{{SALUTATION}}" in msg and "{{BODY}}" in msg
    assert "{{DATE}}" in msg
    assert not out.exists()


@pytest.mark.parametrize("tokens,missing", [
    (("{{BODY}}",), "{{SALUTATION}}"),
    (("{{SALUTATION}}",), "{{BODY}}"),
    ((), "{{SALUTATION}}"),
])
def test_cover_letter_missing_required_placeholder(tmp_path, tokens, missing):
    t = _make_cl_template(tmp_path / "t.docx", tokens=tokens)
    out = tmp_path / "cl.docx"
    with pytest.raises(TemplateError) as exc:
        render_cover_letter(_content(), t, out)
    assert missing in str(exc.value)
    assert not out.exists()


def test_cover_letter_duplicate_placeholder_is_rejected(tmp_path):
    """Two {{BODY}} paragraphs = ambiguous insert point; fail rather than
    guess."""
    t = _make_cl_template(tmp_path / "t.docx", duplicate="{{BODY}}")
    with pytest.raises(TemplateError, match="more than one paragraph"):
        render_cover_letter(_content(), t, tmp_path / "cl.docx")


def test_cover_letter_duplicate_optional_placeholder_is_rejected(tmp_path):
    t = _make_cl_template(tmp_path / "t.docx",
                          tokens=("{{SALUTATION}}", "{{BODY}}", "{{DATE}}"),
                          duplicate="{{DATE}}")
    with pytest.raises(TemplateError, match=r"\{\{DATE\}\}"):
        render_cover_letter(_content(), t, tmp_path / "cl.docx")


@pytest.mark.parametrize("bad", [
    {"salutation": "", "body": ["x"]},
    {"body": ["x"]},
    {"salutation": None, "body": ["x"]},
])
def test_cover_letter_requires_a_salutation(tmp_path, bad):
    t = _make_cl_template(tmp_path / "t.docx")
    with pytest.raises(ValueError, match="salutation"):
        render_cover_letter(bad, t, tmp_path / "cl.docx")


@pytest.mark.parametrize("bad", [
    {"salutation": "Dear X,", "body": []},
    {"salutation": "Dear X,"},
    {"salutation": "Dear X,", "body": None},
])
def test_cover_letter_requires_a_nonempty_body(tmp_path, bad):
    t = _make_cl_template(tmp_path / "t.docx")
    with pytest.raises(ValueError, match="body"):
        render_cover_letter(bad, t, tmp_path / "cl.docx")


def test_cover_letter_validation_runs_before_writing(tmp_path):
    t = _make_cl_template(tmp_path / "t.docx")
    out = tmp_path / "nested" / "cl.docx"
    with pytest.raises(ValueError):
        render_cover_letter({"salutation": "Dear X,"}, t, out)
    assert not out.exists()


# ---------- happy paths ----------

def test_cover_letter_fills_salutation_and_expands_body(tmp_path):
    t = _make_cl_template(tmp_path / "t.docx")
    out = tmp_path / "cl.docx"
    render_cover_letter(_content(), t, out)
    assert _texts(out) == [
        "Static header line", "Dear Hiring Manager,", "First para.", "Second para.",
    ]


def test_cover_letter_removes_the_body_placeholder(tmp_path):
    t = _make_cl_template(tmp_path / "t.docx")
    out = tmp_path / "cl.docx"
    render_cover_letter(_content(), t, out)
    assert "{{BODY}}" not in "\n".join(_texts(out))


def test_cover_letter_leaves_no_raw_token_behind(tmp_path):
    t = _make_cl_template(
        tmp_path / "t.docx",
        tokens=("{{DATE}}", "{{SALUTATION}}", "{{BODY}}", "{{CLOSING}}",
                "{{SIGNOFF_NAME}}"))
    out = tmp_path / "cl.docx"
    render_cover_letter(_content(date="2026-06-06", closing="Sincerely,",
                                signoff_name="Jane Doe"), t, out)
    assert "{{" not in "\n".join(_texts(out))
    assert _texts(out) == [
        "Static header line", "2026-06-06", "Dear Hiring Manager,",
        "First para.", "Second para.", "Sincerely,", "Jane Doe",
    ]


def test_cover_letter_blanks_an_unsupplied_optional_placeholder(tmp_path):
    """A present-but-unfilled token becomes an empty line — never a literal
    {{DATE}} in a letter the user is about to send."""
    t = _make_cl_template(tmp_path / "t.docx",
                          tokens=("{{DATE}}", "{{SALUTATION}}", "{{BODY}}"))
    out = tmp_path / "cl.docx"
    render_cover_letter(_content(), t, out)
    assert _texts(out) == [
        "Static header line", "", "Dear Hiring Manager,",
        "First para.", "Second para.",
    ]


def test_cover_letter_skips_optional_placeholders_absent_from_the_template(tmp_path):
    """The template may carry static text there instead; supplying the field
    must not inject a new paragraph."""
    t = _make_cl_template(tmp_path / "t.docx",
                          static_text=("Static header line", "Sincerely,"))
    out = tmp_path / "cl.docx"
    render_cover_letter(_content(closing="IGNORED", signoff_name="IGNORED"), t, out)
    texts = _texts(out)
    assert "IGNORED" not in texts
    assert texts == ["Static header line", "Sincerely,", "Dear Hiring Manager,",
                     "First para.", "Second para."]


def test_cover_letter_preserves_untouched_static_paragraphs(tmp_path):
    t = _make_cl_template(tmp_path / "t.docx",
                          static_text=("Jane Doe", "jane@example.com", "617-555-0100"))
    out = tmp_path / "cl.docx"
    render_cover_letter(_content(), t, out)
    assert _texts(out)[:3] == ["Jane Doe", "jane@example.com", "617-555-0100"]


def test_cover_letter_body_clones_keep_the_placeholder_formatting(tmp_path):
    """Each body paragraph inherits the {{BODY}} paragraph's style and run
    formatting — that is the whole point of cloning rather than appending."""
    t = _make_cl_template(tmp_path / "t.docx", body_style="CL Body", body_bold=True)
    out = tmp_path / "cl.docx"
    render_cover_letter(_content(), t, out)
    paras = [p for p in Document(str(out)).paragraphs
             if p.text in ("First para.", "Second para.")]
    assert len(paras) == 2
    for p in paras:
        assert p.style.name == "CL Body"
        assert p.runs[0].bold is True


def test_cover_letter_body_order_is_preserved(tmp_path):
    t = _make_cl_template(tmp_path / "t.docx")
    out = tmp_path / "cl.docx"
    render_cover_letter(_content(body=[f"P{i}." for i in range(6)]), t, out)
    assert _texts(out)[1:] == ["Dear Hiring Manager,"] + [f"P{i}." for i in range(6)]


def test_cover_letter_single_body_paragraph(tmp_path):
    t = _make_cl_template(tmp_path / "t.docx")
    out = tmp_path / "cl.docx"
    render_cover_letter(_content(body=["Only one."]), t, out)
    assert _texts(out) == ["Static header line", "Dear Hiring Manager,", "Only one."]


def test_cover_letter_creates_missing_output_dirs(tmp_path):
    t = _make_cl_template(tmp_path / "t.docx")
    out = tmp_path / "applications" / "sap" / "2026-06-06_acme" / "cl.docx"
    render_cover_letter(_content(), t, out)
    assert out.exists()


def test_cover_letter_leaves_the_template_untouched(tmp_path):
    t = _make_cl_template(tmp_path / "t.docx")
    before = t.read_bytes()
    render_cover_letter(_content(), t, tmp_path / "cl.docx")
    assert t.read_bytes() == before


def test_cover_letter_token_split_across_runs_is_still_matched(tmp_path):
    """Word often splits a typed token into several runs; matching on the
    paragraph's full text is what makes a hand-authored template work."""
    doc = Document()
    p = doc.add_paragraph()
    for chunk in ("{{SAL", "UTA", "TION}}"):
        p.add_run(chunk)
    doc.add_paragraph("{{BODY}}")
    t = tmp_path / "t.docx"
    doc.save(str(t))
    out = tmp_path / "cl.docx"
    render_cover_letter(_content(), t, out)
    assert _texts(out) == ["Dear Hiring Manager,", "First para.", "Second para."]


def test_cover_letter_surrounding_whitespace_in_a_token_is_tolerated(tmp_path):
    doc = Document()
    doc.add_paragraph("  {{SALUTATION}}  ")
    doc.add_paragraph("{{BODY}}")
    t = tmp_path / "t.docx"
    doc.save(str(t))
    out = tmp_path / "cl.docx"
    render_cover_letter(_content(), t, out)
    assert "Dear Hiring Manager," in _texts(out)


def test_cover_letter_token_with_other_text_is_not_a_placeholder(tmp_path):
    """Only a paragraph whose ENTIRE text is the token counts — otherwise the
    renderer would mangle prose that happens to mention a token."""
    t = _make_cl_template(tmp_path / "t.docx",
                          static_text=("Use {{DATE}} in your template.",))
    out = tmp_path / "cl.docx"
    render_cover_letter(_content(date="2026-06-06"), t, out)
    assert _texts(out)[0] == "Use {{DATE}} in your template."
