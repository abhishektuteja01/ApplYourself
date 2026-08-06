"""profile-extract feeds /onboarding's resume ingest.

The interesting cases are the ones that silently return nothing: a table-based
resume (python-docx's .paragraphs list is empty for it) and contact details
parked in a header.
"""
from __future__ import annotations

from pathlib import Path

import pytest
from docx import Document

from src.profile_extract import extract, extract_docx, main


def _docx(tmp_path, build) -> Path:
    doc = Document()
    build(doc)
    path = tmp_path / "resume.docx"
    doc.save(str(path))
    return path


class TestDocx:
    def test_extracts_paragraphs_in_order(self, tmp_path):
        def build(doc):
            for line in ("Jane Roe", "Widget Analyst", "Built things."):
                doc.add_paragraph(line)

        text = extract_docx(_docx(tmp_path, build))
        assert text.splitlines() == ["Jane Roe", "Widget Analyst", "Built things."]

    def test_skips_empty_paragraphs(self, tmp_path):
        def build(doc):
            doc.add_paragraph("One")
            doc.add_paragraph("   ")
            doc.add_paragraph("Two")

        assert extract_docx(_docx(tmp_path, build)).splitlines() == ["One", "Two"]

    def test_extracts_table_cells(self, tmp_path):
        """A paragraph-only walk returns nothing for a table-based resume."""

        def build(doc):
            table = doc.add_table(rows=2, cols=2)
            table.cell(0, 0).text = "Employer"
            table.cell(0, 1).text = "2023-2025"
            table.cell(1, 0).text = "Widget Analyst"
            table.cell(1, 1).text = "Exampletown"

        text = extract_docx(_docx(tmp_path, build))
        assert "Employer | 2023-2025" in text
        assert "Widget Analyst | Exampletown" in text

    def test_interleaves_paragraphs_and_tables_in_document_order(self, tmp_path):
        def build(doc):
            doc.add_paragraph("EXPERIENCE")
            table = doc.add_table(rows=1, cols=2)
            table.cell(0, 0).text = "Employer"
            table.cell(0, 1).text = "Dates"
            doc.add_paragraph("EDUCATION")

        lines = extract_docx(_docx(tmp_path, build)).splitlines()
        assert lines == ["EXPERIENCE", "Employer | Dates", "EDUCATION"]

    def test_extracts_header_text(self, tmp_path):
        """Contact details are commonly parked in the header."""

        def build(doc):
            doc.add_paragraph("Body line")
            doc.sections[0].header.paragraphs[0].text = "jane@example.com | 555-0100"

        text = extract_docx(_docx(tmp_path, build))
        assert "[section 1 header]" in text
        assert "jane@example.com | 555-0100" in text

    def test_collapses_merged_cell_duplicates(self, tmp_path):
        def build(doc):
            table = doc.add_table(rows=1, cols=3)
            for cell in table.rows[0].cells:
                cell.text = "Same"

        assert extract_docx(_docx(tmp_path, build)).splitlines() == ["Same"]


class TestPassthroughAndErrors:
    def test_markdown_passes_through(self, tmp_path):
        p = tmp_path / "resume.md"
        p.write_text("**Jane Roe**\n\nBullet\n", encoding="utf-8")
        assert extract(p) == "**Jane Roe**\n\nBullet\n"

    def test_pdf_points_at_the_read_tool(self, tmp_path):
        p = tmp_path / "resume.pdf"
        p.write_bytes(b"%PDF-1.4")
        with pytest.raises(ValueError, match="Read tool"):
            extract(p)

    def test_legacy_doc_explains_the_fix(self, tmp_path):
        p = tmp_path / "resume.doc"
        p.write_bytes(b"\xd0\xcf\x11\xe0")
        with pytest.raises(ValueError, match="save"):
            extract(p)

    def test_unknown_extension_lists_supported_ones(self, tmp_path):
        p = tmp_path / "resume.rtf"
        p.write_text("x", encoding="utf-8")
        with pytest.raises(ValueError, match=r"\.docx"):
            extract(p)

    def test_missing_file_is_an_error(self, tmp_path):
        with pytest.raises(ValueError, match="does not exist"):
            extract(tmp_path / "nope.docx")

    def test_non_utf8_text_is_a_clean_error_not_a_traceback(self, tmp_path):
        """A latin-1 export is common; main() only catches ValueError."""
        p = tmp_path / "resume.md"
        p.write_bytes(b"Jos\xe9 Garc\xeda\n")
        with pytest.raises(ValueError, match="not valid UTF-8"):
            extract(p)


class TestCLI:
    def test_prints_text_and_exits_zero(self, tmp_path, monkeypatch, capsys):
        p = tmp_path / "r.md"
        p.write_text("Jane Roe\n", encoding="utf-8")
        monkeypatch.setattr("sys.argv", ["profile-extract", str(p)])
        assert main() == 0
        assert "Jane Roe" in capsys.readouterr().out

    def test_empty_document_is_an_error_not_silent_success(
        self, tmp_path, monkeypatch, capsys
    ):
        """A scanned-image resume extracts to nothing; saying so beats handing
        the onboarding session an empty string to interview against."""
        p = tmp_path / "r.md"
        p.write_text("   \n", encoding="utf-8")
        monkeypatch.setattr("sys.argv", ["profile-extract", str(p)])
        assert main() == 1
        assert "no text" in capsys.readouterr().err

    def test_unsupported_file_exits_nonzero(self, tmp_path, monkeypatch, capsys):
        p = tmp_path / "r.pdf"
        p.write_bytes(b"%PDF")
        monkeypatch.setattr("sys.argv", ["profile-extract", str(p)])
        assert main() == 1
        assert "ERROR" in capsys.readouterr().err
