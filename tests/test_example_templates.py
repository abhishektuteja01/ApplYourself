"""The two committed .docx templates are the repo's only tracked binaries.

They exist because a Word template cannot be a commented text file, and they are
allowlisted by name in scripts/pii_scan.sh -- which means the text scan that
protects every other tracked file does not protect these two. These tests are
that protection instead:

  - contract: both still satisfy the renderers that consume them, so a fresh
    clone can render a resume and a cover letter without opening Word.
  - guard: neither carries text beyond an approved set, and neither carries an
    author in its core properties. Re-saving either in Word stamps the editor's
    name into docProps/core.xml, which is exactly what the PII gate exists to
    stop and exactly what the allowlist would let through.

The templates are hand-authored in Word, not generated. Run
scripts/scrub_example_templates.py after any Word edit; these tests verify it
was run.
"""
from __future__ import annotations

import importlib.util
import re
import shutil
import zipfile
from pathlib import Path

import pytest
from docx import Document

from src.docx_cover_letter import (
    COVER_LETTER_OPTIONAL_PLACEHOLDERS,
    COVER_LETTER_REQUIRED_PLACEHOLDERS,
    list_cover_letter_placeholders,
    render_cover_letter,
)
from src.docx_render import (
    REQUIRED_CHARACTER_STYLE,
    REQUIRED_STYLES,
    render_resume,
)

REPO_ROOT = Path(__file__).resolve().parent.parent
RESUME_TEMPLATE = REPO_ROOT / "profile" / "resume_template.example.docx"
COVER_TEMPLATE = REPO_ROOT / "profile" / "cover_letter_template.example.docx"
SCRUBBER = REPO_ROOT / "scripts" / "scrub_example_templates.py"

# Hardcoded on purpose: an independent statement of what these files are allowed
# to contain. Every string is a generic stand-in the user replaces in their own
# copy. Adding one means confirming by eye that it carries no personal data.
APPROVED_RESUME_TEXT = {
    ")",
    "Bullet Text",
    "Bullet Text 1",
    "Bullet Text 2",
    "City, ST",
    "College",
    "Course1, Course2, Course3",
    "Date 1 - Present",
    "Date 1 – Date 2",
    "EDUCATION",
    "Email",
    "Master of Science in Computer Science",
    "Name",
    "PROJECTS",
    "Phone",
    "Project 1",
    "SUMMARY",
    "Skill section :",
    "Skill1, Skill2, Skill3",
    "TECHNICAL SKILLS",
    "This is for the summary section",
    "WORK EXPERIENCE",
    "Work 1",
    "www.link.com",
    "x.xx",
    "| Coursework:",
    "| GPA:",
    "•",
    "• [Portfolio](",
}

# The cover letter's letterhead is literal text by design: the user edits their
# own name and contact details into their copy. render_cover_letter preserves
# every non-placeholder paragraph verbatim, so this set is what ships to an
# employer if the user forgets to edit it -- it must stay generic.
# "2" is the footer's page number. "Recipient" is an orphaned glossary building
# block left by the content control scrub_example_templates.py unwraps.
APPROVED_COVER_LETTERHEAD = {
    "2",
    "City, ST",
    "Email",
    "NAME",
    "Num",
    "Recipient",
    "Sincerely,",
    "|",
}
# {{DATE}} is the one optional placeholder the template uses; the closing and
# signoff are literal letterhead instead.
USED_COVER_PLACEHOLDERS = set(COVER_LETTER_REQUIRED_PLACEHOLDERS) | {"{{DATE}}"}

APPROVED_COVER_TEXT = APPROVED_COVER_LETTERHEAD | USED_COVER_PLACEHOLDERS

# The same content seen a paragraph at a time. Word splits a line into several
# runs, so the run-level set above holds "City, ST", "|" and "Email" separately
# while python-docx reports the joined paragraph.
APPROVED_COVER_PARAGRAPHS = {
    "NAME",
    "City, ST | Num| Email",
    "Sincerely,",
} | USED_COVER_PLACEHOLDERS

# Only fields python-docx's CoreProperties actually exposes. "company" and
# "manager" are app.xml properties, not core properties -- asserting them here
# would raise AttributeError, and assigning them in the scrubber would silently
# scrub nothing. They are checked against app.xml instead.
# The run-text element, and ONLY it. A bare <w:t[^>]*> also matches <w:tbl>,
# <w:tc>, <w:tr> and <w:tab/>, which swallows table markup into the "text" and
# makes both the failure messages and the comparison unreliable.
_W_T_RE = re.compile(r"<w:t(?:\s[^>]*)?>(.*?)</w:t>", re.S)

CORE_PROPERTIES_MUST_BE_EMPTY = (
    "author",
    "last_modified_by",
    "title",
    "subject",
    "comments",
    "category",
    "keywords",
)
# Template is included because Word writes the GUID of the personal .dotx the
# document was authored from.
APP_PROPERTIES_MUST_BE_EMPTY = ("Company", "Manager", "HyperlinkBase", "Template")

# Exact part list each template must have. An allowlist, not a denylist: the
# parts that carry identity (word/comments.xml, word/people.xml,
# word/embeddings/*) are exactly the ones nobody thinks to check for.
# The two differ because they were authored from different Word templates.
EXPECTED_RESUME_PARTS = {
    "[Content_Types].xml",
    "_rels/.rels",
    "customXml/_rels/item1.xml.rels",
    "customXml/item1.xml",
    "customXml/itemProps1.xml",
    "docProps/app.xml",
    "docProps/core.xml",
    "word/_rels/document.xml.rels",
    "word/document.xml",
    "word/fontTable.xml",
    "word/numbering.xml",
    "word/settings.xml",
    "word/styles.xml",
    "word/theme/theme1.xml",
    "word/webSettings.xml",
}

# docProps/custom.xml holds only Word's stock-template AssetID (TF10002039), and
# word/glossary/* holds one orphaned "Recipient" building block. Both were
# reviewed and carry no identity; they are listed so a THIRD new part still
# fails this test.
EXPECTED_COVER_PARTS = {
    "[Content_Types].xml",
    "_rels/.rels",
    "docProps/app.xml",
    "docProps/core.xml",
    "docProps/custom.xml",
    "word/_rels/document.xml.rels",
    "word/_rels/settings.xml.rels",
    "word/document.xml",
    "word/endnotes.xml",
    "word/fontTable.xml",
    "word/footer1.xml",
    "word/footnotes.xml",
    "word/glossary/_rels/document.xml.rels",
    "word/glossary/document.xml",
    "word/glossary/fontTable.xml",
    "word/glossary/settings.xml",
    "word/glossary/styles.xml",
    "word/glossary/webSettings.xml",
    "word/header1.xml",
    "word/header2.xml",
    "word/numbering.xml",
    "word/settings.xml",
    "word/styles.xml",
    "word/theme/theme1.xml",
    "word/webSettings.xml",
}


def _load_scrubber():
    spec = importlib.util.spec_from_file_location("scrub_example_templates", SCRUBBER)
    module = importlib.util.module_from_spec(spec)
    spec.loader.exec_module(module)
    return module


def _document_text(path: Path) -> set[str]:
    """Every non-empty body paragraph string, as python-docx sees it.

    Used for the contract tests, which care about what the RENDERERS read.
    Never use it as the PII guard: see _all_package_text.
    """
    doc = Document(str(path))
    return {p.text.strip() for p in doc.paragraphs if p.text.strip()}


def _all_package_text(path: Path) -> set[str]:
    """Every <w:t> string in every XML part of the archive.

    The PII guard has to be closed by construction, and Document.paragraphs is
    not: it reads direct w:r/w:hyperlink children of top-level w:p elements in
    word/document.xml only. Text in a table cell, a content control (w:sdt), a
    tracked insertion (w:ins), a smart tag, a text box (w:txbxContent), a
    footnote, or a comment is invisible to it — and render_cover_letter
    preserves all of those verbatim into a generated letter.
    """
    found: set[str] = set()
    with zipfile.ZipFile(path) as zf:
        for name in zf.namelist():
            if not name.endswith(".xml"):
                continue
            xml = zf.read(name).decode("utf-8")
            found |= {
                text.strip()
                for text in _W_T_RE.findall(xml)
                if text.strip()
            }
    return found


class TestFilesExist:
    def test_both_templates_are_committed(self):
        assert RESUME_TEMPLATE.exists(), f"{RESUME_TEMPLATE} missing"
        assert COVER_TEMPLATE.exists(), f"{COVER_TEMPLATE} missing"

    def test_both_are_valid_zip_archives(self):
        for path in (RESUME_TEMPLATE, COVER_TEMPLATE):
            assert zipfile.is_zipfile(path), f"{path} is not a valid .docx"


class TestResumeTemplateContract:
    def test_declares_every_required_style(self):
        names = {s.name for s in Document(str(RESUME_TEMPLATE)).styles}
        missing = [s for s in REQUIRED_STYLES if s not in names]
        assert not missing, f"template missing paragraph styles: {missing}"
        assert REQUIRED_CHARACTER_STYLE in names

    def test_has_no_tables_or_inline_shapes(self):
        doc = Document(str(RESUME_TEMPLATE))
        assert not doc.tables, "tables are ATS-hostile and the renderer rejects them"
        assert not doc.inline_shapes, "inline shapes are rejected by the renderer"

    def test_renders_a_real_resume(self, tmp_path):
        """The example_primary lane resume is markdown a new clone actually has."""
        md = (
            REPO_ROOT
            / "profile/verticals/example_primary/resume_example_primary.md"
        ).read_text(encoding="utf-8")
        out = tmp_path / "resume.docx"
        render_resume(md, RESUME_TEMPLATE, out)
        doc = Document(str(out))
        used = {p.style.name for p in doc.paragraphs}
        assert used <= set(REQUIRED_STYLES), f"unexpected styles rendered: {used}"
        assert "Resume Bullet" in used and "Resume Name" in used

    def test_demo_text_never_reaches_a_rendered_resume(self, tmp_path):
        """render_resume clears the body, so the style samples are safe to ship.
        If that ever stops being true, the template's demo text starts appearing
        in real resumes."""
        out = tmp_path / "resume.docx"
        # Deliberately shares no string with APPROVED_RESUME_TEXT, or the input
        # itself would trip the assertion below.
        render_resume("**Real Name**\n\nMetropolis, ZZ", RESUME_TEMPLATE, out)
        rendered = _document_text(out)
        assert not (rendered & APPROVED_RESUME_TEXT), (
            f"template demo text leaked into the rendered resume: "
            f"{sorted(rendered & APPROVED_RESUME_TEXT)}"
        )


class TestCoverLetterTemplateContract:
    def test_every_required_placeholder_is_present(self):
        """Asserted through the renderer's own finder, not doc.paragraphs: a
        placeholder inside a content control is invisible to the latter but is
        exactly what the renderer resolves."""
        found = list_cover_letter_placeholders(COVER_TEMPLATE)
        for token in COVER_LETTER_REQUIRED_PLACEHOLDERS:
            assert token in found, f"required placeholder {token} missing"

    def test_declares_no_placeholder_the_renderer_cannot_fill(self):
        """A typo'd token is preserved verbatim, so it ships as literal
        '{{SIGNOFF_NAM}}' inside a real cover letter."""
        declared = set(COVER_LETTER_REQUIRED_PLACEHOLDERS) | set(
            COVER_LETTER_OPTIONAL_PLACEHOLDERS
        )
        stray = {
            token
            for token in re.findall(r"\{\{[A-Z_]+\}\}", " ".join(_all_package_text(COVER_TEMPLATE)))
            if token not in declared
        }
        assert not stray, f"template carries unfillable placeholders: {sorted(stray)}"

    def test_literal_text_is_the_approved_letterhead(self):
        """The letterhead is literal by design -- the user edits their own name
        and contact details into their copy. Every non-placeholder paragraph is
        preserved verbatim into every letter, so it must stay generic."""
        extra = _document_text(COVER_TEMPLATE) - APPROVED_COVER_PARAGRAPHS
        assert not extra, f"template carries unapproved literal text: {sorted(extra)}"

    def test_has_no_tables_or_shapes(self):
        """render_cover_letter preserves everything that is not a placeholder
        paragraph. A letterhead table is the classic way identity text reaches
        every letter while staying invisible to a paragraph-level scan."""
        doc = Document(str(COVER_TEMPLATE))
        assert not doc.tables, "a table here ships verbatim in every letter"
        assert not doc.inline_shapes, "a shape here ships verbatim in every letter"

    def test_headers_and_footers_carry_no_unapproved_text(self):
        """The template has a page-number footer. Headers and footers are
        preserved into every letter, so they are a live leak path -- allowed to
        exist, not allowed to say anything unreviewed."""
        doc = Document(str(COVER_TEMPLATE))
        for i, section in enumerate(doc.sections):
            for label, part in (("header", section.header), ("footer", section.footer)):
                for para in part.paragraphs:
                    text = para.text.strip()
                    if not text:
                        continue
                    assert text in APPROVED_COVER_LETTERHEAD, (
                        f"section {i + 1} {label} carries unapproved text: {text!r}"
                    )

    def test_carries_no_content_controls(self):
        """Word binds content controls to core properties (w:dataBinding). One
        bound to cp:keywords re-syncs its text from a field the scrubber empties,
        silently blanking a placeholder the renderer requires."""
        with zipfile.ZipFile(COVER_TEMPLATE) as zf:
            xml = zf.read("word/document.xml").decode("utf-8")
        assert "<w:sdt>" not in xml and "<w:sdt " not in xml, (
            "content control in document.xml; run scripts/scrub_example_templates.py"
        )
        assert "dataBinding" not in xml

    def test_renders_a_real_cover_letter(self, tmp_path):
        out = tmp_path / "cl.docx"
        render_cover_letter(
            {
                "salutation": "Dear Hiring Manager,",
                "body": ["First paragraph.", "Second paragraph."],
                "date": "January 1, 2026",
            },
            COVER_TEMPLATE,
            out,
        )
        texts = [p.text.strip() for p in Document(str(out)).paragraphs if p.text.strip()]
        assert texts == [
            "NAME",
            "City, ST | Num| Email",
            "January 1, 2026",
            "Dear Hiring Manager,",
            "First paragraph.",
            "Second paragraph.",
            "Sincerely,",
            "NAME",
        ]
        assert not (set(texts) & set(COVER_LETTER_REQUIRED_PLACEHOLDERS)), (
            "an unfilled placeholder survived into the output"
        )

    def test_closing_and_signoff_values_are_dropped_not_rendered(self, tmp_path):
        """The template hardcodes its closing and signoff, so a caller supplying
        closing/signoff_name has those values silently ignored. Pinned so the
        consequence of that design is visible rather than surprising."""
        out = tmp_path / "cl.docx"
        render_cover_letter(
            {
                "salutation": "Dear Hiring Manager,",
                "body": ["Body."],
                "closing": "Warm regards,",
                "signoff_name": "Some Name",
            },
            COVER_TEMPLATE,
            out,
        )
        texts = [p.text.strip() for p in Document(str(out)).paragraphs]
        for dropped in ("Warm regards,", "Some Name"):
            assert dropped not in texts, (
                f"{dropped!r} rendered; the template regained a placeholder for it, "
                f"so this test and the letterhead approval need revisiting"
            )

    def test_an_unsupplied_date_leaves_no_raw_token(self, tmp_path):
        """{{DATE}} is optional. The renderer fills a found placeholder even with
        "", so omitting the date must yield a blank line, never a literal
        '{{DATE}}' in a letter an employer reads."""
        out = tmp_path / "cl.docx"
        render_cover_letter(
            {"salutation": "Dear Hiring Manager,", "body": ["Body."]},
            COVER_TEMPLATE,
            out,
        )
        texts = [p.text for p in Document(str(out)).paragraphs]
        assert not any("{{" in t for t in texts), f"raw token survived: {texts}"


class TestNoPIILeak:
    """The allowlist in pii_scan.sh disables the text scan for these two files.
    These assertions replace it."""

    @pytest.mark.parametrize(
        "path,approved",
        [
            (RESUME_TEMPLATE, APPROVED_RESUME_TEXT),
            (COVER_TEMPLATE, APPROVED_COVER_TEXT),
        ],
        ids=["resume", "cover_letter"],
    )
    def test_text_is_confined_to_the_approved_set(self, path, approved):
        """Scans every XML part, not just body paragraphs — a letterhead table or
        a tracked-change run would otherwise pass this and still ship."""
        extra = _all_package_text(path) - approved
        assert not extra, (
            f"{path.name} gained text outside its approved set: {sorted(extra)}. "
            f"If this is intentional, update the approved set AND confirm the new "
            f"text carries no personal information."
        )

    @pytest.mark.parametrize(
        "path", [RESUME_TEMPLATE, COVER_TEMPLATE], ids=["resume", "cover_letter"]
    )
    def test_carries_no_authored_revisions_or_comments(self, path):
        """w:author appears on tracked changes and comments and carries a real
        name. It is not text, so the scan above cannot see it."""
        with zipfile.ZipFile(path) as zf:
            for name in zf.namelist():
                if not name.endswith(".xml"):
                    continue
                xml = zf.read(name).decode("utf-8")
                assert "w:author" not in xml, f"{path.name}:{name} carries w:author"

    @pytest.mark.parametrize(
        "path,expected",
        [
            (RESUME_TEMPLATE, EXPECTED_RESUME_PARTS),
            (COVER_TEMPLATE, EXPECTED_COVER_PARTS),
        ],
        ids=["resume", "cover_letter"],
    )
    def test_package_parts_are_an_exact_allowlist(self, path, expected):
        """A denylist of known-bad parts cannot hold a file whose only other
        guard is disabled. Any new part — word/comments.xml, word/people.xml,
        word/embeddings/* — must be reviewed deliberately."""
        with zipfile.ZipFile(path) as zf:
            names = set(zf.namelist())
        assert names == expected, (
            f"{path.name} part list changed.\n"
            f"  added:   {sorted(names - expected)}\n"
            f"  removed: {sorted(expected - names)}"
        )

    @pytest.mark.parametrize(
        "path", [RESUME_TEMPLATE, COVER_TEMPLATE], ids=["resume", "cover_letter"]
    )
    def test_core_properties_carry_no_identity(self, path):
        cp = Document(str(path)).core_properties
        for field in CORE_PROPERTIES_MUST_BE_EMPTY:
            value = getattr(cp, field) or ""
            assert not value.strip(), (
                f"{path.name} core property {field!r} is {value!r}. Word stamps "
                f"this on re-save; run "
                f"scripts/scrub_example_templates.py."
            )

    @pytest.mark.parametrize(
        "path", [RESUME_TEMPLATE, COVER_TEMPLATE], ids=["resume", "cover_letter"]
    )
    def test_app_properties_carry_no_identity(self, path):
        with zipfile.ZipFile(path) as zf:
            app = zf.read("docProps/app.xml").decode("utf-8")
        for field in APP_PROPERTIES_MUST_BE_EMPTY:
            match = re.search(rf"<{field}>(.*?)</{field}>", app, re.S)
            value = match.group(1).strip() if match else ""
            assert not value, f"{path.name} app property {field} is {value!r}"

    @pytest.mark.parametrize(
        "path", [RESUME_TEMPLATE, COVER_TEMPLATE], ids=["resume", "cover_letter"]
    )
    def test_no_embedded_images_or_thumbnail(self, path):
        """An image could carry a signature block the paragraph scan above would
        never see."""
        with zipfile.ZipFile(path) as zf:
            names = zf.namelist()
        assert not [n for n in names if n.startswith("word/media/")], (
            f"{path.name} embeds media"
        )
        assert "docProps/thumbnail.jpeg" not in names, (
            f"{path.name} carries an embedded thumbnail image; run "
            f"scripts/scrub_example_templates.py, which strips it"
        )

    @pytest.mark.parametrize(
        "path", [RESUME_TEMPLATE, COVER_TEMPLATE], ids=["resume", "cover_letter"]
    )
    def test_thumbnail_relationship_is_gone_too(self, path):
        """A dangling relationship to a removed part is a malformed package."""
        with zipfile.ZipFile(path) as zf:
            rels = zf.read("_rels/.rels").decode("utf-8")
        assert "thumbnail" not in rels


class TestScrubberStaysInSync:
    """The scrubber is the only thing standing between a Word save and a name in
    a public repo, so it has to actually cover what the guard asserts."""

    @pytest.mark.parametrize(
        "path", [RESUME_TEMPLATE, COVER_TEMPLATE], ids=["resume", "cover_letter"]
    )
    def test_committed_templates_are_already_scrubbed(self, path, tmp_path):
        """The real invariant: running the scrubber on what is committed is a
        no-op. Operates on a copy so a failing run cannot rewrite the tracked
        file."""
        module = _load_scrubber()
        copy = tmp_path / path.name
        shutil.copy(path, copy)
        assert not module.scrub(copy), (
            f"{path.name} still needs scrubbing; run "
            f"uv run python scripts/scrub_example_templates.py"
        )

    def test_scrubber_covers_every_app_property_the_guard_asserts(self):
        module = _load_scrubber()
        missing = set(APP_PROPERTIES_MUST_BE_EMPTY) - set(module.APP_FIELDS)
        assert not missing, f"scrubber does not empty app properties: {sorted(missing)}"

    def test_scrubber_covers_every_core_property_the_guard_asserts(self):
        """python-docx attribute names vs the XML tags the scrubber rewrites."""
        tag_for = {
            "author": "dc:creator",
            "last_modified_by": "cp:lastModifiedBy",
            "title": "dc:title",
            "subject": "dc:subject",
            "comments": "dc:description",
            "category": "cp:category",
            "keywords": "cp:keywords",
        }
        module = _load_scrubber()
        assert set(tag_for) == set(CORE_PROPERTIES_MUST_BE_EMPTY)
        missing = set(tag_for.values()) - set(module.CORE_FIELDS)
        assert not missing, f"scrubber does not empty core fields: {sorted(missing)}"
