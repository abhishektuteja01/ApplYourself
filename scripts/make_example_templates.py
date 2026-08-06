"""Generate the two committed example Word templates.

    uv run python scripts/make_example_templates.py

Regenerates profile/resume_template.example.docx and
profile/cover_letter_template.example.docx from scratch, so the committed
binaries are reproducible rather than opaque.

Contracts these files must satisfy (enforced by tests/test_example_templates.py):
  resume       -- the five paragraph styles in src.docx_render.REQUIRED_STYLES,
                  plus the "Hyperlink" character style, no tables, no inline
                  shapes. Word ships "Hyperlink"; python-docx's default template
                  does NOT, so it is created here.
  cover letter -- every placeholder in
                  src.docx_cover_letter.COVER_LETTER_REQUIRED_PLACEHOLDERS and
                  _OPTIONAL_PLACEHOLDERS, and NOTHING ELSE. Non-placeholder
                  paragraphs are preserved verbatim into every generated letter,
                  so instructional text here would ship inside real cover
                  letters.

Both files carry empty core properties: python-docx stamps author="python-docx",
and Word stamps the real user's name on re-save. A committed .docx with an
author is exactly what the PII gate exists to stop.
"""
from __future__ import annotations

import re
import zipfile
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_LINE_SPACING
from docx.shared import Pt, RGBColor

from src import paths
from src.docx_cover_letter import (
    COVER_LETTER_OPTIONAL_PLACEHOLDERS,
    COVER_LETTER_REQUIRED_PLACEHOLDERS,
)

RESUME_OUT = paths.REPO_ROOT / "profile" / "resume_template.example.docx"
COVER_OUT = paths.REPO_ROOT / "profile" / "cover_letter_template.example.docx"

# ATS-safe: one common sans-serif, single column, no tables, no graphics.
BODY_FONT = "Calibri"
BODY_SIZE = Pt(10.5)

# Demo text in the resume template. render_resume() clears every body paragraph
# before writing, so this text never reaches a generated resume -- it exists so
# the styles are visible and restyleable in Word. Kept in sync with
# tests/test_example_templates.py's approved-text set.
STYLE_DEMOS = [
    ("Resume Name", "YOUR NAME (TEMPLATE)"),
    ("Resume Body", "City, ST | phone | email | portfolio"),
    ("Resume Section", "SECTION HEADER"),
    ("Resume Job Header", "Employer or Project — Title, Dates"),
    ("Resume Body", "Restyle each of these five paragraph styles in Word."),
    ("Resume Bullet", "Bullet text renders in this style."),
]

# Placeholder order in the cover-letter template. Only these strings appear.
COVER_PLACEHOLDER_ORDER = [
    "{{DATE}}",
    "{{SALUTATION}}",
    "{{BODY}}",
    "{{CLOSING}}",
    "{{SIGNOFF_NAME}}",
]


# python-docx's default template inherits an 8KB docProps/thumbnail.jpeg from
# whichever Word document it was authored from. It carries no personal data, but
# it is an image the PII text scan cannot read inside a file the gate allowlists,
# so it is removed rather than shipped unexamined.
THUMBNAIL_PART = "docProps/thumbnail.jpeg"
_THUMBNAIL_REL_RE = re.compile(
    r'<Relationship[^>]*Target="docProps/thumbnail\.jpeg"[^>]*/>'
)


def _scrub_core_properties(doc) -> None:
    """Empty every core property python-docx actually exposes.

    NOT company/manager: those live in docProps/app.xml, not core.xml, and are
    not CoreProperties attributes. Assigning cp.company would silently create an
    unused Python attribute and scrub nothing -- app.xml ships them empty, and
    tests/test_example_templates.py asserts that.
    """
    cp = doc.core_properties
    cp.author = ""
    cp.last_modified_by = ""
    cp.title = ""
    cp.subject = ""
    cp.comments = ""
    cp.category = ""
    cp.keywords = ""


def _strip_thumbnail(path: Path) -> None:
    """Rewrite the archive without the inherited thumbnail, dropping its
    relationship too so the package stays internally consistent."""
    with zipfile.ZipFile(path) as zf:
        entries = [(info, zf.read(info.filename)) for info in zf.infolist()]

    if not any(info.filename == THUMBNAIL_PART for info, _ in entries):
        return

    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as zf:
        for info, data in entries:
            if info.filename == THUMBNAIL_PART:
                continue
            if info.filename == "_rels/.rels":
                rels = _THUMBNAIL_REL_RE.sub("", data.decode("utf-8"))
                data = rels.encode("utf-8")
            zf.writestr(info, data)


def _tighten(fmt, *, before: int = 0, after: int = 2) -> None:
    fmt.space_before = Pt(before)
    fmt.space_after = Pt(after)
    fmt.line_spacing_rule = WD_LINE_SPACING.SINGLE


def _add_paragraph_style(doc, name: str, *, size, bold: bool, base: str = "Normal"):
    style = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
    style.base_style = doc.styles[base]
    style.font.name = BODY_FONT
    style.font.size = size
    style.font.bold = bold
    return style


def build_resume_template(out_path: Path) -> None:
    doc = Document()

    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = BODY_SIZE
    _tighten(normal.paragraph_format)

    name_style = _add_paragraph_style(doc, "Resume Name", size=Pt(16), bold=True)
    _tighten(name_style.paragraph_format, after=0)

    section = _add_paragraph_style(doc, "Resume Section", size=Pt(11), bold=True)
    _tighten(section.paragraph_format, before=8, after=2)

    # Bold comes from inline runs the renderer creates, so the style is not bold.
    job = _add_paragraph_style(doc, "Resume Job Header", size=BODY_SIZE, bold=False)
    _tighten(job.paragraph_format, before=4)

    body = _add_paragraph_style(doc, "Resume Body", size=BODY_SIZE, bold=False)
    _tighten(body.paragraph_format)

    # Based on List Bullet so the glyph and indent come from Word's own
    # numbering definition rather than a hand-rolled one.
    bullet = _add_paragraph_style(
        doc, "Resume Bullet", size=BODY_SIZE, bold=False, base="List Bullet"
    )
    _tighten(bullet.paragraph_format)

    # Word ships "Hyperlink"; python-docx does not. render_resume() rejects a
    # template without it, so create it as a CHARACTER style.
    link = doc.styles.add_style("Hyperlink", WD_STYLE_TYPE.CHARACTER)
    link.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)
    link.font.underline = True

    for style_name, text in STYLE_DEMOS:
        doc.add_paragraph(text, style=style_name)

    _scrub_core_properties(doc)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    _strip_thumbnail(out_path)


def build_cover_letter_template(out_path: Path) -> None:
    doc = Document()
    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(10)

    expected = set(COVER_LETTER_REQUIRED_PLACEHOLDERS) | set(
        COVER_LETTER_OPTIONAL_PLACEHOLDERS
    )
    if set(COVER_PLACEHOLDER_ORDER) != expected:
        raise SystemExit(
            f"placeholder drift: template writes {sorted(COVER_PLACEHOLDER_ORDER)}, "
            f"src.docx_cover_letter declares {sorted(expected)}"
        )

    for token in COVER_PLACEHOLDER_ORDER:
        doc.add_paragraph(token)

    _scrub_core_properties(doc)
    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    _strip_thumbnail(out_path)


def main() -> int:
    build_resume_template(RESUME_OUT)
    build_cover_letter_template(COVER_OUT)
    for p in (RESUME_OUT, COVER_OUT):
        print(f"wrote {p.relative_to(paths.REPO_ROOT)} ({p.stat().st_size} bytes)")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
